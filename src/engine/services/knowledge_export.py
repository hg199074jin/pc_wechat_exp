"""Knowledge Radar — Export cards to Markdown and Word."""
from __future__ import annotations

from datetime import datetime


def _ts_to_str(ts) -> str:
    try:
        return datetime.fromtimestamp(int(ts)).strftime('%Y-%m-%d %H:%M')
    except Exception:
        return ''


def cards_to_markdown(cards: list) -> str:
    """Convert a list of knowledge cards to a Markdown document."""
    parts = ['# 知识沉淀导出', '']
    for card in cards:
        parts.append(f"## {card.get('title', '未命名')}")
        meta = []
        meta.append(f"**评分:** {card.get('score', 0)}")
        meta.append(f"**类型:** {card.get('type', '')}")
        meta.append(f"**状态:** {card.get('status', '')}")
        if card.get('date'):
            meta.append(f"**日期:** {card['date']}")
        if card.get('tags'):
            meta.append(f"**标签:** {', '.join(card['tags'])}")
        parts.append(' | '.join(meta))
        parts.append('')
        if card.get('summary'):
            parts.append(f"> {card['summary']}")
            parts.append('')
        if card.get('why_valuable'):
            parts.append(f"**价值理由:** {card['why_valuable']}")
            parts.append('')
        parts.append(card.get('content_md') or card.get('summary') or '')
        parts.append('')
        sources = card.get('sources') or []
        if sources:
            parts.append('### 来源')
            for src in sources:
                line = f"> [{src.get('chat_name', '')}] {src.get('sender', '')}"
                ts = _ts_to_str(src.get('create_time'))
                if ts:
                    line += f" ({ts})"
                parts.append(line)
                if src.get('quote'):
                    parts.append(f"> {src['quote']}")
                parts.append('')
        parts.append('---')
        parts.append('')
    return '\n'.join(parts)


def cards_to_docx(cards: list) -> bytes:
    """Convert cards to a Word document. Returns bytes or raises ImportError."""
    try:
        from docx import Document
        from docx.shared import Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
    except ImportError:
        raise ImportError('python-docx 未安装，请运行: pip install python-docx')

    doc = Document()
    doc.add_heading('知识沉淀导出', level=0)

    for card in cards:
        doc.add_heading(card.get('title', '未命名'), level=1)

        # Meta line
        meta = doc.add_paragraph()
        meta.add_run(f"评分: {card.get('score', 0)}").bold = True
        meta.add_run(f"  |  类型: {card.get('type', '')}")
        meta.add_run(f"  |  状态: {card.get('status', '')}")
        if card.get('date'):
            meta.add_run(f"  |  日期: {card['date']}")

        if card.get('tags'):
            doc.add_paragraph(f"标签: {', '.join(card['tags'])}")

        if card.get('summary'):
            p = doc.add_paragraph()
            p.style = 'Quote' if 'Quote' in [s.name for s in doc.styles] else None
            p.add_run(card['summary']).italic = True

        if card.get('why_valuable'):
            doc.add_paragraph(f"价值理由: {card['why_valuable']}")

        if card.get('content_md'):
            # Simple: add as plain paragraph (no Markdown rendering in docx)
            for line in card['content_md'].split('\n'):
                if line.strip():
                    doc.add_paragraph(line.strip())

        sources = card.get('sources') or []
        if sources:
            doc.add_heading('来源', level=2)
            for src in sources:
                line = f"[{src.get('chat_name', '')}] {src.get('sender', '')}"
                ts = _ts_to_str(src.get('create_time'))
                if ts:
                    line += f" ({ts})"
                p = doc.add_paragraph()
                p.add_run(line).bold = True
                if src.get('quote'):
                    doc.add_paragraph(src['quote'], style='Quote' if 'Quote' in [s.name for s in doc.styles] else None)

        doc.add_paragraph('')  # spacer

    import io
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
